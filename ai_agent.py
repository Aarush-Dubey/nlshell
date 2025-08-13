#!/usr/bin/env python3
"""
AI Agent Core (ai_agent.py)
Extended AI functionality with automatic command execution and file interpretation.
"""

import os
import json
import mimetypes
import subprocess
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import re
from dataclasses import dataclass

# File type handlers
try:
    import PyPDF2
    import docx
    from PIL import Image
    import pandas as pd
except ImportError:
    pass

@dataclass
class SafetyRule:
    """Safety rule for command execution"""
    pattern: str
    allowed: bool
    reason: str
    auto_execute: bool = False

class AIAgent:
    def __init__(self, ai_core):
        self.ai_core = ai_core
        self.safety_rules = self._init_safety_rules()
        self.file_handlers = self._init_file_handlers()
    
    def _init_safety_rules(self) -> List[SafetyRule]:
        """Initialize safety rules for automatic command execution"""
        return [
            # SAFE - Information gathering commands
            SafetyRule(r'^ls\s', True, "Safe file listing", True),
            SafetyRule(r'^find\s.*-type\s+f.*-exec\s+ls\s+-lh', True, "Safe file search with size", True),
            SafetyRule(r'^du\s+-h.*--max-depth', True, "Safe disk usage check", True),
            SafetyRule(r'^df\s+-h', True, "Safe disk space check", True),
            SafetyRule(r'^wc\s+-l', True, "Safe line count", True),
            SafetyRule(r'^file\s+', True, "Safe file type detection", True),
            SafetyRule(r'^head\s+-n\s+\d+', True, "Safe file preview", True),
            SafetyRule(r'^tail\s+-n\s+\d+', True, "Safe file preview", True),
            SafetyRule(r'^grep\s+.*--color=never', True, "Safe text search", True),
            SafetyRule(r'^stat\s+', True, "Safe file statistics", True),
            SafetyRule(r'^pwd$', True, "Safe current directory", True),
            SafetyRule(r'^whoami$', True, "Safe user info", True),
            SafetyRule(r'^uname\s+-a', True, "Safe system info", True),
            SafetyRule(r'^ps\s+aux', True, "Safe process listing", True),
            SafetyRule(r'^top\s+-b\s+-n1', True, "Safe system snapshot", True),
            SafetyRule(r'^free\s+-h', True, "Safe memory info", True),
            SafetyRule(r'^uptime$', True, "Safe uptime info", True),
            SafetyRule(r'^cat\s+.*\.(txt|log|conf|json|xml|csv)$', True, "Safe text file reading", True),
            
            # RESTRICTED - Potentially dangerous commands
            SafetyRule(r'^rm\s+', False, "File deletion not allowed", False),
            SafetyRule(r'^sudo\s+', False, "Elevated privileges not allowed", False),
            SafetyRule(r'^chmod\s+', False, "Permission changes not allowed", False),
            SafetyRule(r'^chown\s+', False, "Ownership changes not allowed", False),
            SafetyRule(r'^mv\s+', False, "File moving not allowed", False),
            SafetyRule(r'^cp\s+.*>\s*/', False, "System file copying not allowed", False),
            SafetyRule(r'>\s*/etc/', False, "System directory writing not allowed", False),
            SafetyRule(r'>\s*/usr/', False, "System directory writing not allowed", False),
            SafetyRule(r'>\s*/bin/', False, "System directory writing not allowed", False),
            SafetyRule(r'^dd\s+', False, "Direct disk access not allowed", False),
            SafetyRule(r'^fdisk\s+', False, "Disk partitioning not allowed", False),
            SafetyRule(r'^mkfs\s+', False, "Filesystem creation not allowed", False),
            SafetyRule(r'^mount\s+', False, "Mounting not allowed", False),
            SafetyRule(r'^umount\s+', False, "Unmounting not allowed", False),
            SafetyRule(r'^kill\s+-9', False, "Force kill not allowed", False),
            SafetyRule(r'^killall\s+', False, "Mass process termination not allowed", False),
            SafetyRule(r'^shutdown\s+', False, "System shutdown not allowed", False),
            SafetyRule(r'^reboot\s+', False, "System reboot not allowed", False),
            SafetyRule(r'^poweroff\s+', False, "System poweroff not allowed", False),
            SafetyRule(r'^halt\s+', False, "System halt not allowed", False),
            SafetyRule(r'curl.*\|\s*bash', False, "Piped execution not allowed", False),
            SafetyRule(r'wget.*\|\s*bash', False, "Piped execution not allowed", False),
        ]
    
    def _init_file_handlers(self) -> Dict[str, callable]:
        """Initialize file type handlers"""
        return {
            'text/plain': self._handle_text_file,
            'application/json': self._handle_json_file,
            'application/pdf': self._handle_pdf_file,
            'application/msword': self._handle_doc_file,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._handle_docx_file,
            'text/csv': self._handle_csv_file,
            'application/vnd.ms-excel': self._handle_excel_file,
            'image/jpeg': self._handle_image_file,
            'image/png': self._handle_image_file,
            'image/gif': self._handle_image_file,
            'video/mp4': self._handle_video_file,
            'audio/mpeg': self._handle_audio_file,
        }
    
    def check_command_safety(self, command: str) -> Tuple[bool, bool, str]:
        """Check if command is safe to execute
        Returns: (is_safe, auto_execute, reason)
        """
        command_lower = command.lower().strip()
        
        # Check against safety rules
        for rule in self.safety_rules:
            if re.match(rule.pattern, command_lower):
                return rule.allowed, rule.auto_execute and rule.allowed, rule.reason
        
        # Default to safe for basic read-only operations
        safe_patterns = [
            r'^echo\s+',
            r'^which\s+',
            r'^type\s+',
            r'^history\s*$',
            r'^date\s*$',
        ]
        
        for pattern in safe_patterns:
            if re.match(pattern, command_lower):
                return True, True, "Basic safe command"
        
        # Default: not auto-executable but ask user
        return True, False, "Requires user confirmation"
    
    async def auto_execute_safe_command(self, command: str) -> Tuple[int, str, str]:
        """Execute command if it's deemed safe"""
        is_safe, auto_execute, reason = self.check_command_safety(command)
        
        if not is_safe:
            return 1, "", f"Command blocked for safety: {reason}"
        
        if not auto_execute:
            return 1, "", f"Command requires manual confirmation: {reason}"
        
        # Execute the command
        try:
            if command.strip().startswith('cd '):
                path = command.strip()[3:].strip()
                if not path:
                    path = str(Path.home())
                try:
                    os.chdir(os.path.expanduser(path))
                    return 0, f"Changed directory to {os.getcwd()}", ""
                except OSError as e:
                    return 1, "", str(e)
            
            process = await asyncio.create_subprocess_shell(
                command,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                cwd=os.getcwd()
            )
            
            stdout, stderr = await process.communicate()
            return process.returncode, stdout.decode('utf-8'), stderr.decode('utf-8')
            
        except Exception as e:
            return 1, "", str(e)
    
    async def process_agent_query(self, query: str, current_dir: str, history: List[Dict]) -> Dict[str, Any]:
        """Process agent query with automatic command execution and interpretation"""
        
        # Generate commands using AI
        ai_response = await self.ai_core.process_natural_language(query, current_dir, history)
        
        if not ai_response.suggested_commands:
            return {
                'type': 'response',
                'message': ai_response.message,
                'executed_commands': [],
                'results': []
            }
        
        executed_commands = []
        results = []
        final_data = []
        
        # Execute safe commands automatically
        for command in ai_response.suggested_commands:
            is_safe, auto_execute, reason = self.check_command_safety(command)
            
            if auto_execute:
                returncode, stdout, stderr = await self.auto_execute_safe_command(command)
                executed_commands.append(command)
                
                if returncode == 0:
                    results.append({
                        'command': command,
                        'success': True,
                        'output': stdout,
                        'error': ''
                    })
                    final_data.append(stdout)
                else:
                    results.append({
                        'command': command,
                        'success': False,
                        'output': '',
                        'error': stderr
                    })
            else:
                # Command needs user confirmation
                results.append({
                    'command': command,
                    'success': False,
                    'output': '',
                    'error': f"Requires confirmation: {reason}",
                    'needs_confirmation': True
                })
        
        # If we have data, analyze and interpret it
        if final_data:
            interpretation = await self._interpret_command_results(query, executed_commands, final_data)
            
            return {
                'type': 'agent_response',
                'message': ai_response.message,
                'executed_commands': executed_commands,
                'results': results,
                'interpretation': interpretation,
                'auto_executed': True
            }
        else:
            return {
                'type': 'confirmation_needed',
                'message': ai_response.message,
                'suggested_commands': ai_response.suggested_commands,
                'results': results
            }
    
    async def _interpret_command_results(self, original_query: str, commands: List[str], outputs: List[str]) -> str:
        """Interpret command results using AI"""
        
        combined_output = "\n".join(outputs)
        
        prompt = f"""
        TASK: Interpret and analyze command results to answer the user's question.
        
        ORIGINAL QUERY: "{original_query}"
        EXECUTED COMMANDS: {json.dumps(commands)}
        COMMAND OUTPUTS:
        {combined_output}
        
        INSTRUCTIONS:
        1. Analyze the command outputs to directly answer the user's question
        2. Provide specific, actionable insights
        3. If looking for files by size, show the largest files with their sizes
        4. If analyzing disk usage, highlight the directories using most space
        5. Format the response clearly with key findings
        6. Use markdown formatting for better readability
        
        Provide your interpretation:
        """
        
        try:
            interpretation = await self.ai_core._call_ai(prompt)
            return interpretation
        except Exception as e:
            return f"Error interpreting results: {str(e)}"
    
    async def analyze_file(self, filepath: str) -> Dict[str, Any]:
        """Analyze file content based on its type"""
        
        if not os.path.exists(filepath):
            return {
                'error': f"File not found: {filepath}",
                'type': 'error'
            }
        
        # Get file info
        file_stat = os.stat(filepath)
        mime_type, _ = mimetypes.guess_type(filepath)
        
        result = {
            'filepath': filepath,
            'size': file_stat.st_size,
            'mime_type': mime_type,
            'type': 'file_analysis'
        }
        
        # Handle based on file type
        if mime_type in self.file_handlers:
            try:
                content = await self.file_handlers[mime_type](filepath)
                result['content'] = content
            except Exception as e:
                result['error'] = f"Error reading file: {str(e)}"
        else:
            # Try to read as text
            try:
                content = await self._handle_text_file(filepath)
                result['content'] = content
            except:
                result['error'] = "Unsupported file type or binary file"
        
        return result
    
    async def _handle_text_file(self, filepath: str) -> Dict[str, Any]:
        """Handle text files"""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return {
            'type': 'text',
            'preview': content[:2000] + "..." if len(content) > 2000 else content,
            'lines': len(content.splitlines()),
            'characters': len(content)
        }
    
    async def _handle_json_file(self, filepath: str) -> Dict[str, Any]:
        """Handle JSON files"""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return {
            'type': 'json',
            'structure': str(type(data).__name__),
            'keys': list(data.keys()) if isinstance(data, dict) else None,
            'length': len(data) if isinstance(data, (list, dict)) else None,
            'preview': json.dumps(data, indent=2)[:1000] + "..." if len(str(data)) > 1000 else json.dumps(data, indent=2)
        }
    
    async def _handle_pdf_file(self, filepath: str) -> Dict[str, Any]:
        """Handle PDF files"""
        try:
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = len(reader.pages)
                
                # Extract text from first few pages
                text = ""
                for i in range(min(3, pages)):
                    text += reader.pages[i].extract_text()
                
                return {
                    'type': 'pdf',
                    'pages': pages,
                    'text_preview': text[:1500] + "..." if len(text) > 1500 else text
                }
        except ImportError:
            return {'error': 'PyPDF2 not installed. Install with: pip install PyPDF2'}
        except Exception as e:
            return {'error': f'Error reading PDF: {str(e)}'}
    
    async def _handle_docx_file(self, filepath: str) -> Dict[str, Any]:
        """Handle Word documents"""
        try:
            import docx
            doc = docx.Document(filepath)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'type': 'docx',
                'paragraphs': len(doc.paragraphs),
                'text_preview': text[:1500] + "..." if len(text) > 1500 else text
            }
        except ImportError:
            return {'error': 'python-docx not installed. Install with: pip install python-docx'}
        except Exception as e:
            return {'error': f'Error reading DOCX: {str(e)}'}
    
    async def _handle_csv_file(self, filepath: str) -> Dict[str, Any]:
        """Handle CSV files"""
        try:
            import pandas as pd
            df = pd.read_csv(filepath)
            
            return {
                'type': 'csv',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'preview': df.head().to_string(),
                'data_types': df.dtypes.to_dict()
            }
        except ImportError:
            # Fallback without pandas
            with open(filepath, 'r') as f:
                lines = f.readlines()
                return {
                    'type': 'csv',
                    'rows': len(lines) - 1,  # Subtract header
                    'preview': ''.join(lines[:10])
                }
        except Exception as e:
            return {'error': f'Error reading CSV: {str(e)}'}
    
    async def _handle_image_file(self, filepath: str) -> Dict[str, Any]:
        """Handle image files"""
        try:
            from PIL import Image
            with Image.open(filepath) as img:
                return {
                    'type': 'image',
                    'dimensions': f"{img.width}x{img.height}",
                    'format': img.format,
                    'mode': img.mode,
                    'size_mb': round(os.path.getsize(filepath) / (1024*1024), 2)
                }
        except ImportError:
            return {'error': 'Pillow not installed. Install with: pip install Pillow'}
        except Exception as e:
            return {'error': f'Error reading image: {str(e)}'}
    
    async def _handle_video_file(self, filepath: str) -> Dict[str, Any]:
        """Handle video files"""
        return {
            'type': 'video',
            'size_mb': round(os.path.getsize(filepath) / (1024*1024), 2),
            'note': 'Video analysis requires additional tools like ffprobe'
        }
    
    async def _handle_audio_file(self, filepath: str) -> Dict[str, Any]:
        """Handle audio files"""
        return {
            'type': 'audio',
            'size_mb': round(os.path.getsize(filepath) / (1024*1024), 2),
            'note': 'Audio analysis requires additional tools like mutagen'
        }
    
    async def _handle_doc_file(self, filepath: str) -> Dict[str, Any]:
        """Handle old Word documents (.doc)"""
        return {
            'type': 'doc',
            'error': 'Legacy .doc files require additional tools. Consider converting to .docx'
        }
    
    async def _handle_excel_file(self, filepath: str) -> Dict[str, Any]:
        """Handle Excel files"""
        try:
            import pandas as pd
            df = pd.read_excel(filepath)
            
            return {
                'type': 'excel',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'preview': df.head().to_string()
            }
        except ImportError:
            return {'error': 'openpyxl not installed. Install with: pip install openpyxl'}
        except Exception as e:
            return {'error': f'Error reading Excel: {str(e)}'}